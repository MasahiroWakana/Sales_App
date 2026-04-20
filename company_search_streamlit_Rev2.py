import streamlit as st
import json
import os
import time
from datetime import datetime
import requests
from bs4 import BeautifulSoup
import urllib.parse
import re
import csv
import io
import zipfile
import sys
from pathlib import Path
from openpyxl import load_workbook
from docx import Document
import tempfile

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

APP_BASE_DIR = Path(__file__).resolve().parent
CONFIG_FILE = str(APP_BASE_DIR / "company_analyzer_config_streamlit.json")
NTA_CACHE_DIR = APP_BASE_DIR / "nta_houjin_cache"
TEMP_UPLOAD_DIR = APP_BASE_DIR / "streamlit_uploads"
TEMP_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


class SerpAPILimitError(Exception):
    pass


def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def save_config(cfg):
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def get_streamlit_secret(key: str, default: str = "") -> str:
    try:
        value = st.secrets.get(key, default)
        if value is None:
            return default
        return str(value)
    except Exception:
        return default


REQUIRED_LICENSES = {
    "建設": ["建設業許可", "JAC加入"],
    "運送": ["一般貨物運送許可", "運行管理者"],
    "外食": ["飲食店営業許可", "食品衛生責任者"],
    "宿泊": ["旅館業許可"],
    "ビルクリーニング": ["建築物清掃業登録"],
    "製造": ["工場稼働許可", "安全体制"],
    "介護": ["介護事業指定"],
}


def get_required_licenses_list(guessed_industry: str):
    if not guessed_industry:
        return []
    matched = []
    for industry, licenses in REQUIRED_LICENSES.items():
        if industry in guessed_industry:
            matched.extend(licenses)
    seen = set()
    uniq = []
    for lic in matched:
        if lic not in seen:
            uniq.append(lic)
            seen.add(lic)
    return uniq


def get_required_licenses_text(guessed_industry: str) -> str:
    licenses = get_required_licenses_list(guessed_industry)
    return "、".join(licenses) if licenses else "該当なし"


def build_license_queries(company_name: str, guessed_industry: str):
    licenses = get_required_licenses_list(guessed_industry)
    return [f"{company_name} {lic}" for lic in licenses]


def normalize_duckduckgo_url(url: str) -> str:
    if url.startswith("//"):
        url = "https:" + url
    parsed = urllib.parse.urlparse(url)
    qs = urllib.parse.parse_qs(parsed.query)
    if "uddg" in qs:
        return urllib.parse.unquote(qs["uddg"][0])
    return url


def duckduckgo_search_html(query, max_results=5, log_func=None):
    if log_func:
        log_func(f"[DDG] DuckDuckGo(HTML)で検索中: {query}")
    url = "https://duckduckgo.com/html/"
    params = {"q": query, "kl": "jp-ja"}
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
    try:
        r = requests.get(url, params=params, headers=headers, timeout=20)
        r.raise_for_status()
    except Exception as e:
        if log_func:
            log_func(f"[ERROR] DuckDuckGo HTML検索失敗: {e}")
        return []
    soup = BeautifulSoup(r.text, "html.parser")
    results = []
    for a in soup.select("a.result__a"):
        raw_url = a.get("href")
        title = a.get_text(strip=True)
        if not raw_url or not title:
            continue
        clean_url = normalize_duckduckgo_url(raw_url)
        results.append({"title": title, "url": clean_url})
        if len(results) >= max_results:
            break
    if log_func:
        log_func(f"[DDG] HTML検索 結果件数: {len(results)} 件")
    return results


def fetch_url_text(url, max_chars=8000, log_func=None):
    url = normalize_duckduckgo_url(url)
    if log_func:
        log_func(f"[取得] {url}")
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
    try:
        r = requests.get(url, headers=headers, timeout=25)
        r.raise_for_status()
    except Exception as e:
        if log_func:
            log_func(f"[ERROR] 取得失敗: {e}")
        return ""
    content_type = r.headers.get("Content-Type", "").lower()
    if "pdf" in content_type:
        if log_func:
            log_func(f"[INFO] PDFコンテンツ検出のため本文抽出をスキップ: {url}")
        return ""
    soup = BeautifulSoup(r.text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [ln.strip() for ln in text.splitlines()]
    text = "\n".join(ln for ln in lines if ln)
    return text[:max_chars]


def search_via_serpapi_duckduckgo(query, serpapi_key, max_results=5, log_func=None):
    if log_func:
        log_func(f"[SerpAPI] DuckDuckGoエンジンで検索中: {query}")
    endpoint = "https://serpapi.com/search.json"
    params = {"engine": "duckduckgo", "q": query, "api_key": serpapi_key}
    try:
        r = requests.get(endpoint, params=params, timeout=25)
    except Exception as e:
        if log_func:
            log_func(f"[ERROR] SerpAPI 接続失敗: {e}")
        raise SerpAPILimitError(str(e))
    if r.status_code in (401, 402, 403, 429):
        raise SerpAPILimitError(f"SerpAPI status={r.status_code}")
    if r.status_code != 200:
        raise SerpAPILimitError(f"SerpAPI status={r.status_code}")
    data = r.json()
    results = []
    for item in data.get("organic_results", []):
        title = item.get("title")
        link = item.get("link") or item.get("href")
        if title and link:
            results.append({"title": title, "url": link})
        if len(results) >= max_results:
            break
    return results


def fetch_wikipedia_company(company_name, lang="ja", log_func=None):
    search_endpoint = f"https://{lang}.wikipedia.org/w/api.php"
    search_params = {"action": "query", "list": "search", "srsearch": company_name, "format": "json", "srlimit": 3}
    try:
        r = requests.get(search_endpoint, params=search_params, timeout=15)
        r.raise_for_status()
    except Exception:
        return None
    data = r.json()
    search_results = data.get("query", {}).get("search", [])
    if not search_results:
        return None
    page = search_results[0]
    pageid = page.get("pageid")
    title = page.get("title")
    extract_params = {"action": "query", "prop": "extracts", "explaintext": 1, "pageids": pageid, "format": "json"}
    try:
        r2 = requests.get(search_endpoint, params=extract_params, timeout=15)
        r2.raise_for_status()
    except Exception:
        return None
    data2 = r2.json()
    pageinfo = data2.get("query", {}).get("pages", {}).get(str(pageid), {})
    summary = pageinfo.get("extract", "")
    return {"title": title, "url": f"https://{lang}.wikipedia.org/?curid={pageid}", "summary": summary[:8000]}


def create_openai_client(api_key):
    if OpenAI is None:
        raise RuntimeError("openai ライブラリがインストールされていません。pip install openai")
    if not api_key:
        raise RuntimeError("OpenAI API Key が設定されていません。")
    return OpenAI(api_key=api_key)


def openai_call_text(api_key, system_prompt, user_prompt, model_name="gpt-4.1-mini"):
    client = create_openai_client(api_key)
    resp = client.responses.create(
        model=model_name,
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    try:
        return resp.output[0].content[0].text
    except Exception:
        return str(resp)


def plan_search_queries(company_name, api_key, log_func=None, model_name="gpt-4.1-mini"):
    system_prompt = "あなたは日本の企業情報のリサーチを行うアナリストです。与えられた企業名から、日本語で適切な検索クエリを作成してください。"
    user_prompt = f'''企業名: {company_name}
以下の4項目を JSON 形式で返してください（余計な説明文は不要）:
{{
  "official_site_query": "...",
  "kanpo_query": "...",
  "industry_stats_query": "...",
  "guessed_industry": "..."
}}'''
    text = openai_call_text(api_key, system_prompt, user_prompt, model_name=model_name)
    try:
        data = json.loads(text)
    except Exception:
        data = {
            "official_site_query": f"{company_name} 公式サイト 会社概要",
            "kanpo_query": f"{company_name} 決算公告 官報",
            "industry_stats_query": f"{company_name} 業界 売上高 統計",
            "guessed_industry": "不明",
        }
    data.setdefault("official_site_query", f"{company_name} 公式サイト 会社概要")
    data.setdefault("kanpo_query", f"{company_name} 決算公告 官報")
    data.setdefault("industry_stats_query", f"{company_name} 業界 売上高 統計")
    data.setdefault("guessed_industry", "不明")
    return data


def apply_email_domain_to_queries(plan, email_domain: str):
    if email_domain and email_domain.strip():
        base = email_domain.strip()
        plan["official_site_query"] += f" {base}"
        plan["kanpo_query"] += f" {base}"
        plan["industry_stats_query"] += f" {base}"
    return plan


def score_official_site_candidate(title: str, url: str, email_domain: str = "") -> int:
    score = 0
    url_lower = url.lower()
    for kw in ["会社概要", "企業情報", "corporate", "about", "会社情報"]:
        if kw in title:
            score += 3
    if ".co.jp" in url_lower:
        score += 5
    elif ".jp" in url_lower:
        score += 3
    elif ".com" in url_lower:
        score += 2
    if "blog" in url_lower or "news" in url_lower:
        score -= 2
    if email_domain:
        base = email_domain.lstrip("@").lower()
        if base and base in url_lower:
            score += 30
    return score


def collect_sources(company_name, homepage_url, api_key, cfg, log_func=None, model_name="gpt-4.1-mini"):
    email_domain = cfg.get("email_domain", "")
    plan = plan_search_queries(company_name, api_key, log_func, model_name=model_name)
    plan = apply_email_domain_to_queries(plan, email_domain)
    use_serpapi = cfg.get("use_serpapi", True)
    serpapi_key = cfg.get("serpapi_key")
    sources = []
    added_urls = set()

    if homepage_url:
        homepage_text = fetch_url_text(homepage_url, log_func=log_func)
        if homepage_text.strip():
            sources.append({"category": "official", "title": f"{company_name} 公式ホームページ（入力URL）", "url": homepage_url, "text": homepage_text})
            added_urls.add(normalize_duckduckgo_url(homepage_url))

    official_candidates = []
    official_query = plan["official_site_query"]
    if use_serpapi and serpapi_key:
        try:
            serp_results = search_via_serpapi_duckduckgo(official_query, serpapi_key, max_results=8, log_func=log_func)
            for item in serp_results:
                official_candidates.append((score_official_site_candidate(item["title"], item["url"], email_domain), item))
        except SerpAPILimitError:
            cfg["use_serpapi"] = False
            save_config(cfg)
            use_serpapi = False
    if not official_candidates:
        ddg_results = duckduckgo_search_html(official_query, max_results=8, log_func=log_func)
        for item in ddg_results:
            official_candidates.append((score_official_site_candidate(item["title"], item["url"], email_domain), item))
    official_candidates.sort(key=lambda x: x[0], reverse=True)
    for _, item in official_candidates[:3]:
        normalized = normalize_duckduckgo_url(item["url"])
        if normalized in added_urls:
            continue
        text = fetch_url_text(item["url"], log_func=log_func)
        if text.strip():
            sources.append({"category": "official", "title": item["title"], "url": item["url"], "text": text})
            added_urls.add(normalized)

    for category, query in [("kanpo", plan["kanpo_query"]), ("industry", plan["industry_stats_query"])]:
        query_results = []
        if use_serpapi and serpapi_key:
            try:
                query_results = search_via_serpapi_duckduckgo(query, serpapi_key, max_results=5, log_func=log_func)
            except SerpAPILimitError:
                cfg["use_serpapi"] = False
                save_config(cfg)
                use_serpapi = False
        if not query_results:
            query_results = duckduckgo_search_html(query, max_results=5, log_func=log_func)
        for item in query_results[:5]:
            normalized = normalize_duckduckgo_url(item["url"])
            if normalized in added_urls:
                continue
            text = fetch_url_text(item["url"], log_func=log_func)
            if text.strip():
                sources.append({"category": category, "title": item["title"], "url": item["url"], "text": text})
                added_urls.add(normalized)

    wiki = fetch_wikipedia_company(company_name, lang="ja", log_func=log_func)
    if wiki and wiki.get("summary") and wiki["url"] not in added_urls:
        sources.append({"category": "wikipedia", "title": wiki["title"], "url": wiki["url"], "text": wiki["summary"]})
        added_urls.add(wiki["url"])

    guessed_industry = plan.get("guessed_industry", "不明")
    for lic_query in build_license_queries(company_name, guessed_industry):
        lic_results = []
        if use_serpapi and serpapi_key:
            try:
                lic_results = search_via_serpapi_duckduckgo(lic_query, serpapi_key, max_results=3, log_func=log_func)
            except SerpAPILimitError:
                cfg["use_serpapi"] = False
                save_config(cfg)
                use_serpapi = False
        if not lic_results:
            lic_results = duckduckgo_search_html(lic_query, max_results=3, log_func=log_func)
        for item in lic_results[:3]:
            normalized = normalize_duckduckgo_url(item["url"])
            if normalized in added_urls:
                continue
            text = fetch_url_text(item["url"], log_func=log_func)
            if text.strip():
                sources.append({"category": "license", "title": item["title"], "url": item["url"], "text": text})
                added_urls.add(normalized)

    return plan, sources


def build_markdown_report(company_name, homepage_url, plan, sources, api_key, log_func=None, model_name="gpt-4.1-mini"):
    chunks = []
    for i, s in enumerate(sources, start=1):
        header = f"### Source {i} ({s['category']})\nタイトル: {s['title']}\nURL: {s['url']}\n\n"
        body = s["text"][:6000]
        chunks.append(header + body)
    combined_sources = "\n\n".join(chunks) if chunks else "（ソーステキストなし）"
    guessed_industry = plan.get("guessed_industry", "不明")
    required_licenses_text = get_required_licenses_text(guessed_industry)
    system_prompt = "あなたは日本の企業情報を収集・要約するシニアアナリストです。根拠が薄い場合はその旨を明記し、無理な断定は避けてください。"
    user_prompt = f'''対象企業名: {company_name}
指定ホームページURL: {homepage_url}
想定業種: {guessed_industry}
確認対象の必須資格・許認可: {required_licenses_text}

以下に企業公式サイト・官報・統計・業界統計・Wikipediaと思われるソースを示します。
これらを総合的に読み、次の要件を満たす Markdown レポートを日本語で作成してください。

# 要件
1. サマリー
2. 主要指標（表形式）
3. 経営指標(売上高,経常利益,利益余剰金)（表形式）
4. 必須資格・許認可の確認（表形式）
5. 詳細メモ
6. 注意点

以下がソース全文です:

{combined_sources}
'''
    return openai_call_text(api_key, system_prompt, user_prompt, model_name=model_name)


JP_PREFECTURES = [
    "北海道", "青森県", "岩手県", "宮城県", "秋田県", "山形県", "福島県",
    "茨城県", "栃木県", "群馬県", "埼玉県", "千葉県", "東京都", "神奈川県",
    "新潟県", "富山県", "石川県", "福井県", "山梨県", "長野県",
    "岐阜県", "静岡県", "愛知県", "三重県", "滋賀県", "京都府", "大阪府", "兵庫県", "奈良県", "和歌山県",
    "鳥取県", "島根県", "岡山県", "広島県", "山口県", "徳島県", "香川県", "愛媛県", "高知県",
    "福岡県", "佐賀県", "長崎県", "熊本県", "大分県", "宮崎県", "鹿児島県", "沖縄県",
]


def guess_head_office_address_from_text(text: str) -> str:
    if not text:
        return ""
    compact = re.sub(r"\s+", " ", text.replace("　", " "))
    patterns = [
        r"(?:本社所在地|本店所在地|所在地|住所|本社)[:：\s]*((?:北海道|東京都|京都府|大阪府|..県).{5,80})",
        r"((?:北海道|東京都|京都府|大阪府|..県)(?:[^\n]{5,100}?)(?:\d{1,4}[-−ー‐]\d{1,4}(?:[-−ー‐]\d{1,4})?))",
    ]
    for pat in patterns:
        m = re.search(pat, compact)
        if m:
            addr = m.group(1).strip(" 　:：,，。")
            for stop in ["電話", "TEL", "FAX", "代表", "アクセス", "地図", "MAP"]:
                pos = addr.find(stop)
                if pos > 0:
                    addr = addr[:pos].strip()
            return addr[:120]
    return ""


def extract_identity_candidates_from_sources(company_name: str, homepage_url: str, sources: list):
    candidates = []
    if homepage_url:
        candidates.append({"title": f"{company_name} 入力URL", "url": homepage_url, "text": fetch_url_text(homepage_url, max_chars=12000)})
    for s in sources:
        if s.get("category") == "official":
            candidates.append({"title": s.get("title", ""), "url": s.get("url", ""), "text": s.get("text", "")[:12000]})
    uniq = []
    seen = set()
    for c in candidates:
        key = c.get("url", "") + "|" + c.get("title", "")
        if key not in seen:
            uniq.append(c)
            seen.add(key)
        if len(uniq) >= 4:
            break
    return uniq


def analyze_company_identity(company_name: str, homepage_url: str, sources: list, api_key: str, log_func=None, model_name="gpt-4.1-mini"):
    candidates = extract_identity_candidates_from_sources(company_name, homepage_url, sources)
    fallback_address = ""
    text_blocks = []
    for i, c in enumerate(candidates, start=1):
        txt = c.get("text", "")
        if txt and not fallback_address:
            fallback_address = guess_head_office_address_from_text(txt)
        text_blocks.append(f"### Candidate {i}\nタイトル: {c.get('title', '')}\nURL: {c.get('url', '')}\n\n{txt[:8000]}")
    joined = "\n\n".join(text_blocks)
    if not joined:
        return {"official_company_name": company_name, "head_office_address": "", "evidence": "公式ソース本文を取得できなかったため特定不能", "confidence": "低"}
    system_prompt = "あなたは日本企業の基本情報を公開情報から抽出するアナリストです。明示されている内容だけを使ってください。"
    user_prompt = f'''対象企業名: {company_name}
入力URL: {homepage_url}
以下の公式サイト候補本文から、対象企業の正式名称と本社住所を抽出してください。
必ず JSON のみで返してください。
{{
  "official_company_name": "...",
  "head_office_address": "...",
  "evidence": "本文中の根拠要約",
  "confidence": "高/中/低"
}}
本文:
{joined}'''
    try:
        raw = openai_call_text(api_key, system_prompt, user_prompt, model_name=model_name)
        data = json.loads(raw)
    except Exception:
        data = {"official_company_name": company_name, "head_office_address": fallback_address, "evidence": "AI抽出に失敗したため、公式サイト本文から機械的に住所候補を抽出", "confidence": "低"}
    return {
        "official_company_name": (data.get("official_company_name") or company_name).strip(),
        "head_office_address": (data.get("head_office_address") or "").strip() or fallback_address,
        "evidence": data.get("evidence", ""),
        "confidence": data.get("confidence", "低"),
    }


def get_prefecture_city_for_nta(address: str):
    address = (address or "").strip()
    prefecture, city = "", ""
    for pref in JP_PREFECTURES:
        if address.startswith(pref):
            prefecture = pref
            rest = address[len(pref):]
            m = re.match(r"(.+?(?:市|区|町|村|郡))", rest)
            if m:
                city = m.group(1)
            break
    return prefecture, city


def try_parse_nta_number_from_text(text: str):
    if not text:
        return []
    items = []
    for m in re.finditer(r"(\d{13}).{0,120}?((?:北海道|東京都|京都府|大阪府|..県)[^\n]{3,160})", text):
        items.append({"corporate_number": m.group(1), "address": m.group(2).strip(), "name": ""})
    return items


def ensure_nta_cache_dir():
    NTA_CACHE_DIR.mkdir(parents=True, exist_ok=True)


def normalize_address_for_nta_match(address: str) -> str:
    s = (address or "").strip().replace("　", "")
    s = re.sub(r"\s+", "", s)
    s = s.replace("丁目", "-").replace("番地", "-").replace("番", "-").replace("号", "")
    s = s.replace("−", "-").replace("ー", "-").replace("―", "-").replace("‐", "-")
    s = re.sub(r"-+", "-", s)
    return s.strip("-")


def normalize_company_name_for_match(value: str) -> str:
    if value is None:
        return ""
    s = str(value).strip()
    s = s.replace("株式会社", "").replace("有限会社", "").replace("合同会社", "")
    s = s.replace("（株）", "").replace("(株)", "")
    s = s.replace("　", "").replace(" ", "")
    s = re.sub(r"[\-ー‐―−・,，\.\(\)（）\[\]【】]", "", s)
    return s.lower()


def normalize_company_name_for_nta_search(name: str) -> str:
    s = normalize_company_name_for_match(name)
    for token in ["株式会社", "有限会社", "合同会社", "合名会社", "合資会社", "（株）", "㈱", "(株)"]:
        s = s.replace(token, "")
    return s


def get_nta_download_page_html(log_func=None):
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)", "Referer": "https://www.houjin-bangou.nta.go.jp/"}
    urls = ["https://www.houjin-bangou.nta.go.jp/download/zenken/", "https://www.houjin-bangou.nta.go.jp/download/zenken/index.html"]
    for url in urls:
        try:
            r = requests.get(url, headers=headers, timeout=30)
            if r.status_code == 200 and r.text:
                r.encoding = r.apparent_encoding or r.encoding
                return r.text, url
        except Exception as e:
            if log_func:
                log_func(f"[WARN] 国税庁ダウンロードページ取得失敗: {e}")
    return "", urls[0]


def _extract_unicode_section_html(page_html: str) -> str:
    if not page_html:
        return ""
    m = re.search(r'CSV形式[^<]{0,20}Unicode.*?(?=XML形式[^<]{0,20}Unicode)', page_html, flags=re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(0)
    m = re.search(r'CSV.*?Unicode.*?(?=XML.*?Unicode)', page_html, flags=re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(0)
    return page_html


def _find_zip_urls_near_keyword(section_html: str, keyword: str, base_url: str, window: int = 2500):
    urls = []
    if not section_html or not keyword:
        return urls
    for m in re.finditer(re.escape(keyword), section_html):
        snippet = section_html[m.start():m.start() + window]
        hrefs = re.findall(r'href=["\']([^"\']+?\.zip(?:\?[^"\']*)?)["\']', snippet, flags=re.IGNORECASE)
        for href in hrefs:
            full = urllib.parse.urljoin(base_url, href)
            if full not in urls:
                urls.append(full)
        if urls:
            return urls
    return urls


def _find_prefecture_zip_urls(prefecture: str, page_html: str, base_url: str, log_func=None):
    section_html = _extract_unicode_section_html(page_html)
    urls = _find_zip_urls_near_keyword(section_html, prefecture, base_url)
    if urls:
        return urls
    urls = _find_zip_urls_near_keyword(page_html, prefecture, base_url)
    if urls:
        return urls
    soup = BeautifulSoup(page_html, "html.parser")
    urls = []
    for a in soup.find_all("a", href=True):
        href = urllib.parse.urljoin(base_url, a["href"])
        if ".zip" not in href.lower():
            continue
        context = " ".join([a.get_text(" ", strip=True), a.parent.get_text(" ", strip=True) if a.parent else ""])
        if prefecture in context and href not in urls:
            urls.append(href)
    if urls:
        return urls
    fallback_urls = _find_zip_urls_near_keyword(section_html or page_html, "全国", base_url)
    return fallback_urls[:1]


def _download_nta_zip_url(zip_url: str, log_func=None):
    ensure_nta_cache_dir()
    parsed = urllib.parse.urlparse(zip_url)
    filename = os.path.basename(parsed.path) or "nta_download.zip"
    if parsed.query:
        stem, ext = os.path.splitext(filename)
        filename = f"{stem}_{str(abs(hash(parsed.query)))[:8]}{ext or '.zip'}"
    local_path = NTA_CACHE_DIR / filename
    if local_path.exists() and local_path.stat().st_size > 0:
        return str(local_path)
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)", "Referer": "https://www.houjin-bangou.nta.go.jp/download/zenken/"}
    try:
        r = requests.get(zip_url, headers=headers, timeout=180)
        r.raise_for_status()
        local_path.write_bytes(r.content)
        return str(local_path)
    except Exception as e:
        if log_func:
            log_func(f"[WARN] 国税庁ZIPダウンロード失敗: {e} / {zip_url}")
        return ""


def download_nta_prefecture_zips(prefecture: str, log_func=None):
    if not prefecture:
        return []
    html, base_url = get_nta_download_page_html(log_func=log_func)
    zip_urls = _find_prefecture_zip_urls(prefecture, html, base_url, log_func=log_func)
    local_paths = []
    for zip_url in zip_urls:
        path = _download_nta_zip_url(zip_url, log_func=log_func)
        if path:
            local_paths.append(path)
    return local_paths


def _open_all_csv_rows_from_zip(zip_path: str):
    with zipfile.ZipFile(zip_path, 'r') as zf:
        names = [n for n in zf.namelist() if n.lower().endswith('.csv')]
        rows = []
        for csv_name in names:
            raw = zf.read(csv_name)
            text = None
            for enc in ('utf-8-sig', 'utf-8', 'cp932', 'shift_jis', 'latin1'):
                try:
                    text = raw.decode(enc)
                    break
                except Exception:
                    continue
            if text is None:
                continue
            for row in csv.reader(io.StringIO(text)):
                if row:
                    rows.append(row)
        return rows


def _open_nta_csv_rows(csv_path: str):
    rows = []
    with open(csv_path, 'r', encoding='shift_jis', errors='replace', newline='') as f:
        for row in csv.reader(f):
            if row:
                rows.append(row)
    return rows


def _extract_nta_row_candidate(row):
    if not row or len(row) < 4:
        return None
    corp_idx = -1
    corp_no = ""
    for idx, val in enumerate(row[:12]):
        vv = (val or '').strip().replace('"', '')
        if re.fullmatch(r'\d{13}', vv):
            corp_idx = idx
            corp_no = vv
            break
    if corp_idx < 0:
        return None
    clean = [(c or '').strip() for c in row]
    name = ""
    for j in range(corp_idx + 1, min(len(clean), corp_idx + 10)):
        v = clean[j]
        if v and not re.fullmatch(r'[0-9\-/.:]+', v):
            if any(ch in v for ch in ['株式会社', '有限会社', '合同会社', '協同組合', '運送', '物流']) or re.search(r'[一-龥ぁ-んァ-ヴ]', v):
                name = v
                break
    addr_start = -1
    for j in range(corp_idx + 1, min(len(clean), corp_idx + 15)):
        v = clean[j]
        if any(v.startswith(pref) for pref in JP_PREFECTURES):
            addr_start = j
            break
    address = ""
    if addr_start >= 0:
        parts = []
        for j in range(addr_start, min(len(clean), addr_start + 4)):
            v = clean[j]
            if not v:
                continue
            parts.append(v)
            joined = ''.join(parts)
            if re.search(r'\d{1,4}[-−ー‐]\d{1,4}', joined) or '丁目' in joined or '番地' in joined:
                address = joined
                break
        if not address:
            address = ''.join(parts)
    if not address:
        joined = ''.join(clean[max(0, corp_idx - 1):min(len(clean), corp_idx + 15)])
        m = re.search(r'(北海道|東京都|京都府|大阪府|..県).{3,120}', joined)
        if m:
            address = m.group(0)
    return {"corporate_number": corp_no, "name": name, "address": address}


def search_corporate_number_from_nta_csv(company_name: str, head_office_address: str, csv_path: str, log_func=None):
    pref, city = get_prefecture_city_for_nta(head_office_address)
    if not csv_path or not os.path.exists(csv_path):
        return {"searched_name": company_name, "searched_address": head_office_address, "prefecture": pref, "city": city, "candidates": [], "status": "未取得", "source": "法人番号CSV"}
    rows = _open_nta_csv_rows(csv_path)
    search_name_norm = normalize_company_name_for_nta_search(company_name)
    search_addr_norm = normalize_address_for_nta_match(head_office_address)
    city_norm = normalize_address_for_nta_match(city)
    candidates = []
    for row in rows:
        item = _extract_nta_row_candidate(row)
        if not item:
            continue
        item_name_norm = normalize_company_name_for_nta_search(item.get('name', ''))
        item_addr_norm = normalize_address_for_nta_match(item.get('address', ''))
        score = 0
        if search_name_norm and item_name_norm:
            if search_name_norm == item_name_norm:
                score += 120
            elif search_name_norm in item_name_norm or item_name_norm in search_name_norm:
                score += 80
        if search_addr_norm and item_addr_norm:
            if search_addr_norm == item_addr_norm:
                score += 120
            elif search_addr_norm in item_addr_norm or item_addr_norm in search_addr_norm:
                score += 80
        if city_norm and item_addr_norm and city_norm in item_addr_norm:
            score += 20
        if pref and item.get('address', '').startswith(pref):
            score += 10
        if score >= 80:
            item['score'] = score
            candidates.append(item)
    uniq = []
    seen = set()
    for item in sorted(candidates, key=lambda x: x.get('score', 0), reverse=True):
        key = (item.get('corporate_number'), item.get('name'), item.get('address'))
        if key in seen:
            continue
        uniq.append(item)
        seen.add(key)
    return {"searched_name": company_name, "searched_address": head_office_address, "prefecture": pref, "city": city, "candidates": uniq, "status": "候補あり" if uniq else "未取得", "source": "法人番号CSV"}


def search_corporate_number_from_nta_download(company_name: str, head_office_address: str, log_func=None):
    pref, city = get_prefecture_city_for_nta(head_office_address)
    if not pref:
        return {"searched_name": company_name, "searched_address": head_office_address, "prefecture": pref, "city": city, "candidates": [], "status": "未取得", "source": "国税庁法人番号公表サイト 全件データ"}
    zip_paths = download_nta_prefecture_zips(pref, log_func=log_func)
    search_name_norm = normalize_company_name_for_nta_search(company_name)
    search_addr_norm = normalize_address_for_nta_match(head_office_address)
    city_norm = normalize_address_for_nta_match(city)
    candidates = []
    for zip_path in zip_paths:
        try:
            rows = _open_all_csv_rows_from_zip(zip_path)
        except Exception:
            continue
        for row in rows:
            item = _extract_nta_row_candidate(row)
            if not item:
                continue
            item_name_norm = normalize_company_name_for_nta_search(item.get('name', ''))
            item_addr_norm = normalize_address_for_nta_match(item.get('address', ''))
            score = 0
            if search_name_norm and item_name_norm:
                if search_name_norm == item_name_norm:
                    score += 120
                elif search_name_norm in item_name_norm or item_name_norm in search_name_norm:
                    score += 80
            if search_addr_norm and item_addr_norm:
                if search_addr_norm == item_addr_norm:
                    score += 120
                elif search_addr_norm in item_addr_norm or item_addr_norm in search_addr_norm:
                    score += 80
            if city_norm and item_addr_norm and city_norm in item_addr_norm:
                score += 20
            if pref and item.get('address', '').startswith(pref):
                score += 10
            if score >= 80:
                item['score'] = score
                candidates.append(item)
    uniq = []
    seen = set()
    for item in sorted(candidates, key=lambda x: x.get('score', 0), reverse=True):
        key = (item.get('corporate_number', ''), item.get('name', ''), item.get('address', ''))
        if key in seen:
            continue
        seen.add(key)
        uniq.append(item)
        if len(uniq) >= 20:
            break
    return {"searched_name": company_name, "searched_address": head_office_address, "prefecture": pref, "city": city, "candidates": uniq, "status": "候補あり" if uniq else "未取得", "source": "国税庁法人番号公表サイト 全件データ"}


def search_corporate_number_from_nta_web(company_name: str, head_office_address: str, log_func=None):
    candidates = []
    pref, city = get_prefecture_city_for_nta(head_office_address)
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
    trial_requests = [
        ("https://www.houjin-bangou.nta.go.jp/", {}),
        ("https://www.houjin-bangou.nta.go.jp/download/zenken/", {}),
        ("https://www.houjin-bangou.nta.go.jp/ks/", {"q": company_name}),
        ("https://www.houjin-bangou.nta.go.jp/ks/", {"q": company_name, "pref": pref, "city": city}),
        ("https://www.houjin-bangou.nta.go.jp/ks/", {"q": company_name, "address": head_office_address}),
    ]
    for url, params in trial_requests:
        params = {k: v for k, v in params.items() if v}
        try:
            r = requests.get(url, params=params, headers=headers, timeout=25)
            if r.status_code != 200:
                continue
            parsed = try_parse_nta_number_from_text(r.text)
            if parsed:
                candidates.extend(parsed)
                break
        except Exception:
            pass
    return {"searched_name": company_name, "searched_address": head_office_address, "prefecture": pref, "city": city, "candidates": candidates, "status": "候補あり" if candidates else "未取得", "source": "国税庁法人番号公表サイト Webページ"}


def search_corporate_number_from_nta(company_name: str, head_office_address: str, log_func=None, csv_path: str = ""):
    csv_result = None
    if csv_path:
        csv_result = search_corporate_number_from_nta_csv(company_name, head_office_address, csv_path, log_func=log_func)
        if csv_result.get('candidates'):
            return csv_result
    dl_result = search_corporate_number_from_nta_download(company_name, head_office_address, log_func=log_func)
    if dl_result.get('candidates'):
        return dl_result
    web_result = search_corporate_number_from_nta_web(company_name, head_office_address, log_func=log_func)
    if web_result.get('candidates'):
        return web_result
    return csv_result if csv_path else dl_result


def choose_best_corporate_number_candidate(company_name: str, head_office_address: str, nta_result: dict):
    address_norm = normalize_address_for_nta_match(head_office_address)
    name_norm = normalize_company_name_for_nta_search(company_name)
    best, best_score, scored_candidates = None, -1, []
    for item in nta_result.get("candidates", []):
        score = int(item.get("score", 0) or 0)
        item = dict(item)
        addr_norm = normalize_address_for_nta_match(item.get("address", ""))
        item_name_norm = normalize_company_name_for_nta_search(item.get("name", ""))
        if name_norm and item_name_norm:
            if name_norm == item_name_norm:
                score += 60
            elif name_norm in item_name_norm or item_name_norm in name_norm:
                score += 30
        if address_norm and addr_norm:
            if address_norm == addr_norm:
                score += 80
            elif address_norm in addr_norm or addr_norm in address_norm:
                score += 50
            else:
                pref1, city1 = get_prefecture_city_for_nta(head_office_address)
                pref2, city2 = get_prefecture_city_for_nta(item.get("address", ""))
                if pref1 and pref1 == pref2:
                    score += 15
                if city1 and city1 == city2:
                    score += 15
        item["score"] = score
        scored_candidates.append(item)
        if score > best_score:
            best_score = score
            best = dict(item)
    return {"status": "特定" if best else "不明", "best": best, "candidates": sorted(scored_candidates, key=lambda x: x.get('score', 0), reverse=True), "source": nta_result.get("source", "国税庁法人番号公表サイト")}


def search_general_freight_permit_by_corporate_number(company_name: str, corporate_number: str, cfg: dict, log_func=None):
    use_serpapi = cfg.get("use_serpapi", True)
    serpapi_key = cfg.get("serpapi_key")
    queries = []
    if corporate_number:
        queries.extend([
            f'"{corporate_number}" "一般貨物自動車運送事業"',
            f'"{corporate_number}" "一般貨物自動車運送事業許可"',
            f'"{corporate_number}" "貨物自動車運送事業法"',
        ])
    queries.extend([f'{company_name} 一般貨物自動車運送事業 許可', f'{company_name} 一般貨物自動車運送事業者'])
    results, seen = [], set()
    for query in queries:
        query_results = []
        if use_serpapi and serpapi_key:
            try:
                query_results = search_via_serpapi_duckduckgo(query, serpapi_key=serpapi_key, max_results=5, log_func=log_func)
            except SerpAPILimitError:
                cfg["use_serpapi"] = False
                save_config(cfg)
                use_serpapi = False
        if not query_results:
            query_results = duckduckgo_search_html(query, max_results=5, log_func=log_func)
        for item in query_results:
            url = normalize_duckduckgo_url(item["url"])
            if url in seen:
                continue
            text = fetch_url_text(url, log_func=log_func)
            results.append({"query": query, "title": item["title"], "url": url, "text": text})
            seen.add(url)
            if len(results) >= 10:
                return results
    return results


def summarize_excel_row(ws, row_idx: int, max_cols: int = 8) -> str:
    values = []
    max_col = min(ws.max_column, max_cols)
    for col in range(1, max_col + 1):
        v = ws.cell(row=row_idx, column=col).value
        if v is not None:
            sv = str(v).strip()
            if sv:
                values.append(sv)
    return " / ".join(values)[:300]


def search_company_in_excel(excel_path: str, company_name: str, log_func=None):
    result = {"status": "未実施", "path": excel_path or "", "matches": [], "error": ""}
    if not excel_path:
        result["status"] = "ファイル未指定"
        return result
    if not os.path.exists(excel_path):
        result["status"] = "ファイルなし"
        result["error"] = "指定ファイルが存在しません。"
        return result
    company_norm = normalize_company_name_for_match(company_name)
    if not company_norm:
        result["status"] = "検索不可"
        result["error"] = "企業名が空です。"
        return result
    try:
        wb = load_workbook(excel_path, data_only=True, read_only=True)
        exact_hits, partial_hits = [], []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    val = cell.value
                    if val is None:
                        continue
                    cell_text = str(val).strip()
                    if not cell_text:
                        continue
                    cell_norm = normalize_company_name_for_match(cell_text)
                    if not cell_norm:
                        continue
                    hit = False
                    score = 0
                    if company_norm == cell_norm:
                        hit, score = True, 100
                    elif company_norm in cell_norm or cell_norm in company_norm:
                        hit, score = True, 70
                    if hit:
                        item = {"sheet": ws.title, "cell": cell.coordinate, "value": cell_text[:200], "row": cell.row, "row_summary": summarize_excel_row(ws, cell.row), "score": score}
                        (exact_hits if score >= 100 else partial_hits).append(item)
        matches = exact_hits + partial_hits
        result["matches"] = matches[:10]
        result["status"] = "登録あり" if matches else "登録なし"
        return result
    except Exception as e:
        result["status"] = "読込エラー"
        result["error"] = str(e)
        return result


def evaluate_general_freight_permit(company_name: str, permit_sources: list, log_func=None):
    company_norm = normalize_company_name_for_match(company_name)
    positive_keywords = ["一般貨物自動車運送事業", "一般貨物自動車運送事業許可", "一般貨物自動車運送事業者", "一般貨物運送許可"]
    caution_keywords = ["貨物軽自動車運送事業", "第一種貨物利用運送事業", "第二種貨物利用運送事業"]
    matched_sources = []
    caution_sources = []
    for src in permit_sources:
        merged = "\n".join([str(src.get("title", "")), str(src.get("url", "")), str(src.get("text", ""))[:5000]])
        merged_norm = normalize_company_name_for_match(merged)
        has_company = company_norm and company_norm in merged_norm
        pos = any(k in merged for k in positive_keywords)
        caution = any(k in merged for k in caution_keywords)
        if has_company and pos:
            matched_sources.append({"title": src.get("title", ""), "url": src.get("url", ""), "query": src.get("query", ""), "evidence": next((k for k in positive_keywords if k in merged), "")})
        if has_company and caution:
            caution_sources.append({"title": src.get("title", ""), "url": src.get("url", ""), "query": src.get("query", ""), "evidence": next((k for k in caution_keywords if k in merged), "")})
    status = "あり" if matched_sources else "不明"
    note = "公開情報上で、企業名と一般貨物自動車運送事業に関する記述が同時に確認できました。" if matched_sources else "公開情報から明確な許可記載を確認できませんでした。未取得を意味するものではありません。"
    if caution_sources:
        note += " なお、一般貨物以外の運送類型に関する記述も一部検出されています。"
    return {"status": status, "matches": matched_sources[:5], "cautions": caution_sources[:5], "note": note, "searched_count": len(permit_sources)}


def build_additional_registry_section(company_name: str, identity_result: dict, corporate_number_result: dict, permit_result: dict, gmark_result: dict, workplace_result: dict) -> str:
    lines = []
    lines.append("\n\n---\n")
    lines.append("## 追加調査結果（企業同定 / 法人番号 / 一般貨物自動車運送事業許可 / 認証照合）\n")
    lines.append(f"対象企業: **{company_name}**\n")
    official_name = identity_result.get("official_company_name", "")
    head_office_address = identity_result.get("head_office_address", "")
    corp_best = corporate_number_result.get("best") or {}
    corp_no = corp_best.get("corporate_number", "")
    lines.append("\n### 1. 企業同定結果\n")
    lines.append("|項目|内容|備考|")
    lines.append("|---|---|---|")
    lines.append(f"|特定企業名|{official_name or '不明'}|入力企業名と公式サイト本文から抽出|")
    lines.append(f"|本社住所|{head_office_address or '不明'}|{identity_result.get('evidence', '')}|")
    lines.append(f"|同定信頼度|{identity_result.get('confidence', '低')}|公式サイト本文ベース|")
    lines.append("\n### 2. 法人番号確認（国税庁法人番号公表サイト）\n")
    lines.append("|項目|内容|備考|")
    lines.append("|---|---|---|")
    lines.append(f"|検索状態|{corporate_number_result.get('status', '不明')}|取得元: {corporate_number_result.get('source', '国税庁法人番号公表サイト')}|")
    lines.append(f"|法人番号|{corp_no or '不明'}|国税庁サイト候補照合結果|")
    lines.append(f"|国税庁側所在地|{corp_best.get('address', '') or '不明'}|対象本社住所との照合スコア: {corp_best.get('score', '') if corp_best else ''}|")
    if corporate_number_result.get("candidates"):
        lines.append("\n|No|法人番号|所在地|スコア|")
        lines.append("|---|---|---|---|")
        for i, c in enumerate(corporate_number_result["candidates"][:10], start=1):
            lines.append(f"|{i}|{c.get('corporate_number', '')}|{c.get('address', '')}|{c.get('score', '')}|")
    else:
        lines.append("\n国税庁法人番号公表サイトから一致候補を取得できませんでした。")
    lines.append("\n### 3. 一般貨物自動車運送事業の許可確認（追加）\n")
    lines.append(f"- 判定: **{permit_result.get('status', '不明')}**")
    lines.append(f"- コメント: {permit_result.get('note', '')}")
    lines.append(f"- 使用した法人番号: {corp_no or '未取得'}")
    lines.append(f"- 検索で確認した候補件数: {permit_result.get('searched_count', 0)} 件\n")
    if permit_result.get("matches"):
        lines.append("|No|検索クエリ|タイトル|URL|確認語句|")
        lines.append("|---|---|---|---|---|")
        for i, m in enumerate(permit_result["matches"], start=1):
            lines.append(f"|{i}|{m.get('query', '')}|{m.get('title', '')}|{m.get('url', '')}|{m.get('evidence', '')}|")
    else:
        lines.append("公開情報から、法人番号または企業名と一般貨物自動車運送事業の記述が同時に確認できる明示ソースは見つかりませんでした。")
    lines.append("\n### 4. 認証照合\n")
    lines.append("|確認項目|判定|根拠|備考|")
    lines.append("|---|---|---|---|")
    gmark_root = gmark_result.get("path", "") or "未指定"
    if gmark_result.get("matches"):
        hit = gmark_result["matches"][0]
        gmark_root = f"{os.path.basename(gmark_result.get('path', ''))} / {hit.get('sheet', '')}!{hit.get('cell', '')}"
    lines.append(f"|Gマーク認証|{gmark_result.get('status', '未実施')}|{gmark_root}|Excel照合結果|")
    work_root = workplace_result.get("path", "") or "未指定"
    if workplace_result.get("matches"):
        hit = workplace_result["matches"][0]
        work_root = f"{os.path.basename(workplace_result.get('path', ''))} / {hit.get('sheet', '')}!{hit.get('cell', '')}"
    lines.append(f"|働きやすい職場認証|{workplace_result.get('status', '未実施')}|{work_root}|Excel照合結果|")
    return "\n".join(lines)




def _clean_md_cell(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip())


def add_markdown_table_to_docx(doc: Document, table_lines: list[str]):
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if cells:
            rows.append(cells)
    if len(rows) < 2:
        for line in table_lines:
            doc.add_paragraph(line)
        return
    header = rows[0]
    body = []
    for row in rows[1:]:
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in row):
            continue
        body.append(row)
    col_count = max(len(header), max((len(r) for r in body), default=len(header)))
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    for i in range(col_count):
        table.rows[0].cells[i].text = _clean_md_cell(header[i] if i < len(header) else "")
    for row in body:
        cells = table.add_row().cells
        for i in range(col_count):
            cells[i].text = _clean_md_cell(row[i] if i < len(row) else "")


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    doc = Document()
    lines = (markdown_text or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_markdown_table_to_docx(doc, block)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped == "---":
            doc.add_paragraph("―" * 20)
            i += 1
            continue
        if stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            body = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(body, style="List Number")
            i += 1
            continue
        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                break
            if nxt.startswith(("#", "|", "- ", "* ")) or re.match(r"^\d+\.\s+", nxt) or nxt == "---":
                break
            paragraph_lines.append(nxt)
            i += 1
        doc.add_paragraph("\n".join(paragraph_lines))
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def persist_uploaded_file(uploaded_file, suffix=None):
    if uploaded_file is None:
        return ""
    filename = uploaded_file.name
    if suffix and not filename.lower().endswith(suffix.lower()):
        filename = filename + suffix
    safe_name = re.sub(r'[^A-Za-z0-9._-ぁ-んァ-ヴ一-龥]', '_', filename)
    out_path = TEMP_UPLOAD_DIR / f"{int(time.time() * 1000)}_{safe_name}"
    out_path.write_bytes(uploaded_file.getbuffer())
    return str(out_path)


def resolve_input_path(label, path_text, uploaded_file):
    # Windowsローカル実行なら C:\... を優先。クラウド等では uploader を使う。
    if path_text and os.path.exists(path_text):
        return path_text
    if uploaded_file is not None:
        return persist_uploaded_file(uploaded_file)
    return path_text or ""



# =========================================================
# 企業資本関係調査（Streamlit統合版）
# =========================================================
REL_MAX_RESULT_URLS_TO_OPEN = 5
REL_MAX_PAGE_TEXT_LEN = 20000
REL_TRUSTED_PATTERNS = [
    "kabutan.jp",
    "prtimes.jp",
    "nikkei.com",
    "reuters.com",
    "tdnet",
    "edinet-fsa.go.jp",
    "disclosure.edinet-fsa.go.jp",
]
REL_GROUP_LINK_KEYWORDS = [
    "グループ", "関連会社", "グループ会社", "会社情報", "企業情報", "会社概要", "about", "company"
]
REL_CORP_SUFFIXES = ["株式会社", "有限会社", "合同会社", "合名会社", "合資会社"]
REL_COMPANY_NAME_PATTERNS = [
    r"株式会社[^\s\u3000,，、。()（）<>\[\]{}]{1,50}",
    r"有限会社[^\s\u3000,，、。()（）<>\[\]{}]{1,50}",
    r"合同会社[^\s\u3000,，、。()（）<>\[\]{}]{1,50}",
    r"[^\s\u3000,，、。()（）<>\[\]{}]{1,50}株式会社",
    r"[^\s\u3000,，、。()（）<>\[\]{}]{1,50}有限会社",
    r"[^\s\u3000,，、。()（）<>\[\]{}]{1,50}ホールディングス株式会社",
    r"[^\s\u3000,，、。()（）<>\[\]{}]{1,50}ホールディングス",
    r"[^\s\u3000,，、。()（）<>\[\]{}]{1,50}HD",
]
REL_STRONG_SIGNAL_WORDS = [
    "子会社化", "株式取得", "完全子会社", "出資", "資本提携", "M&A", "グループ化", "グループ会社",
    "親会社", "子会社", "傘下", "連結子会社", "持分法適用関連会社", "主要株主", "100%取得", "100％取得",
    "株式100%", "株式100％", "グループの一員",
]


def rel_normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def rel_normalize_company_name(name: str) -> str:
    text = rel_normalize_text(name)
    for suffix in REL_CORP_SUFFIXES:
        text = text.replace(suffix, "")
    text = text.replace("ホールディングス", "HD")
    text = text.replace("・", "")
    text = text.replace("（株）", "").replace("(株)", "")
    return text.strip().lower()


def rel_is_similar_name(a: str, b: str, threshold: float = 0.84) -> bool:
    from difflib import SequenceMatcher
    a_n = rel_normalize_company_name(a)
    b_n = rel_normalize_company_name(b)
    if not a_n or not b_n:
        return False
    if a_n == b_n:
        return True
    return SequenceMatcher(None, a_n, b_n).ratio() >= threshold


def rel_unique_by_name(items: list[dict]) -> list[dict]:
    result, seen = [], []
    for item in items:
        name = rel_normalize_text(item.get("name", ""))
        if not name:
            continue
        if any(rel_is_similar_name(name, s) for s in seen):
            continue
        seen.append(name)
        item = dict(item)
        item["name"] = name
        result.append(item)
    return result


def rel_unique_by_url(items: list[dict]) -> list[dict]:
    result, seen = [], set()
    for item in items:
        url = (item.get("url") or "").strip()
        if not url or url in seen:
            continue
        seen.add(url)
        result.append(dict(item))
    return result


def rel_source_score(url: str, homepage_url: str = "") -> int:
    u = (url or "").lower()
    homepage_netloc = urllib.parse.urlparse(homepage_url).netloc.lower() if homepage_url else ""
    if homepage_netloc and homepage_netloc in u:
        return 100
    if "edinet-fsa.go.jp" in u or "disclosure.edinet-fsa.go.jp" in u:
        return 95
    if "nikkei.com" in u:
        return 85
    if "reuters.com" in u:
        return 80
    if "kabutan.jp" in u or "prtimes.jp" in u or "tdnet" in u:
        return 75
    if ".co.jp" in u:
        return 70
    return 30


def rel_confidence_from_score(score: int, strong_signal: bool = False) -> str:
    if score >= 95 or (score >= 85 and strong_signal):
        return "高"
    if score >= 70:
        return "中"
    return "低"


def rel_domain_label(url: str) -> str:
    try:
        return urllib.parse.urlparse(url).netloc
    except Exception:
        return "unknown"


def rel_extract_company_names(text: str) -> list[str]:
    text = rel_normalize_text(text)
    results = []
    for pattern in REL_COMPANY_NAME_PATTERNS:
        for match in re.findall(pattern, text):
            name = rel_normalize_text(match)
            if 2 <= len(name) <= 60:
                results.append(name)
    deduped = []
    for name in results:
        if not any(rel_is_similar_name(name, s) for s in deduped):
            deduped.append(name)
    return deduped


def rel_contains_strong_signal(text: str) -> bool:
    norm = rel_normalize_text(text)
    return any(word in norm for word in REL_STRONG_SIGNAL_WORDS)


def rel_search_web(query: str, cfg: dict, log_func=None, max_results: int = 8) -> list[dict]:
    use_serpapi = cfg.get("use_serpapi", True)
    serpapi_key = cfg.get("serpapi_key", "")
    if use_serpapi and serpapi_key:
        try:
            return search_via_serpapi_duckduckgo(query, serpapi_key=serpapi_key, max_results=max_results, log_func=log_func)
        except Exception:
            pass
    return duckduckgo_search_html(query, max_results=max_results, log_func=log_func)


def rel_guess_company_url(company_name: str, cfg: dict, log_func=None) -> str:
    queries = [f"{company_name} 公式サイト", f"{company_name} 会社概要", f"{company_name} コーポレートサイト"]
    for query in queries:
        for item in rel_search_web(query, cfg, log_func=log_func, max_results=8):
            url = item.get("url", "")
            title = item.get("title", "")
            if not url or "houjin.info" in url:
                continue
            if company_name in title or rel_normalize_company_name(company_name) in rel_normalize_company_name(title):
                return normalize_duckduckgo_url(url)
    return ""


def rel_find_candidate_detail_pages(company_url: str, log_func=None) -> list[str]:
    try:
        r = requests.get(company_url, headers={"User-Agent": "Mozilla/5.0"}, timeout=25)
        r.raise_for_status()
        raw_html = r.text
    except Exception:
        return [company_url] if company_url else []
    soup = BeautifulSoup(raw_html, "html.parser")
    candidates = [company_url]
    for a in soup.find_all("a", href=True):
        text = rel_normalize_text(a.get_text(" ", strip=True))
        href = (a.get("href") or "").strip()
        if not href:
            continue
        full_url = urllib.parse.urljoin(company_url, href)
        if any(k.lower() in text.lower() for k in REL_GROUP_LINK_KEYWORDS):
            candidates.append(full_url)
    deduped, seen = [], set()
    for u in candidates:
        nu = normalize_duckduckgo_url(u)
        if nu not in seen:
            seen.add(nu)
            deduped.append(nu)
    return deduped[:8]


def rel_extract_relation_names_from_page(text: str, base_url: str) -> tuple[list[dict], list[dict]]:
    subsidiaries, affiliates = [], []
    lines = re.split(r"[\n\r]|。", text)
    for line in lines:
        line_n = rel_normalize_text(line)
        if not line_n:
            continue
        names = rel_extract_company_names(line_n)
        for name in names:
            if "子会社" in line_n or "グループ会社" in line_n or "連結子会社" in line_n:
                subsidiaries.append({"name": name, "note": f"公式サイト本文から抽出: {base_url}", "url": base_url, "source_type": "official", "confidence": "中"})
            elif "関連会社" in line_n or "持分法" in line_n or "パートナー" in line_n:
                affiliates.append({"name": name, "note": f"公式サイト本文から抽出: {base_url}", "url": base_url, "source_type": "official", "confidence": "中"})
    return rel_unique_by_name(subsidiaries), rel_unique_by_name(affiliates)


def rel_scrape_official_site(company_url: str, log_func=None) -> dict:
    subsidiaries, affiliates = [], []
    pages = rel_find_candidate_detail_pages(company_url, log_func=log_func)
    for page_url in pages:
        text = fetch_url_text(page_url, max_chars=REL_MAX_PAGE_TEXT_LEN, log_func=log_func)
        if not text:
            continue
        s_items, a_items = rel_extract_relation_names_from_page(text, page_url)
        subsidiaries.extend(s_items)
        affiliates.extend(a_items)
    return {"subsidiaries": rel_unique_by_name(subsidiaries), "affiliates": rel_unique_by_name(affiliates), "official_pages_checked": pages}


def rel_extract_parent_candidate_from_title(company_name: str, title: str):
    normalized = rel_normalize_text(title)
    if not rel_contains_strong_signal(normalized):
        return None
    patterns = [
        rf"^(.*?)、{re.escape(company_name)}.*?(?:子会社化|株式取得|完全子会社|出資|資本提携|M&A|グループ化)",
        rf"^(.*?)が{re.escape(company_name)}.*?(?:子会社化|株式取得|完全子会社|出資|資本提携|M&A|グループ化)",
        rf"^(.*?)は{re.escape(company_name)}.*?(?:子会社化|株式取得|完全子会社|出資|資本提携|M&A|グループ化)",
    ]
    for pattern in patterns:
        m = re.search(pattern, normalized)
        if m:
            candidate = rel_normalize_text(m.group(1))
            if 1 < len(candidate) <= 80:
                return candidate
    names = rel_extract_company_names(normalized)
    for name in names:
        if not rel_is_similar_name(name, company_name):
            return name
    return None


def rel_build_parent_patterns(company_name: str) -> list:
    cn = re.escape(company_name)
    return [
        re.compile(rf"([\w\-・&＆\.\s一-龥ぁ-んァ-ヶ]+?(?:ホールディングス株式会社|ホールディングス|HD|株式会社))が{cn}の株式100[%％]を取得"),
        re.compile(rf"{cn}は([\w\-・&＆\.\s一-龥ぁ-んァ-ヶ]+?(?:ホールディングス株式会社|ホールディングス|HD|株式会社))の子会社"),
        re.compile(rf"{cn}は([\w\-・&＆\.\s一-龥ぁ-んァ-ヶ]+?(?:ホールディングス株式会社|ホールディングス|HD|株式会社))傘下"),
        re.compile(rf"{cn}は([\w\-・&＆\.\s一-龥ぁ-んァ-ヶ]+?)グループの一員"),
        re.compile(rf"([\w\-・&＆\.\s一-龥ぁ-んァ-ヶ]+?(?:ホールディングス株式会社|ホールディングス|HD|株式会社))は{cn}を(?:完全子会社化|子会社化|グループ化)"),
        re.compile(rf"親会社[：:\s]+([\w\-・&＆\.\s一-龥ぁ-んァ-ヶ]+?(?:ホールディングス株式会社|ホールディングス|HD|株式会社))"),
        re.compile(rf"{cn}.*?主要株主[：:\s]+([\w\-・&＆\.\s一-龥ぁ-んァ-ヶ]+?(?:ホールディングス株式会社|ホールディングス|HD|株式会社))"),
    ]


def rel_extract_parent_from_page_content(company_name: str, url: str, text: str, homepage_url: str = ""):
    patterns = rel_build_parent_patterns(company_name)
    strong_signal = rel_contains_strong_signal(text)
    for pattern in patterns:
        m = pattern.search(text)
        if not m:
            continue
        parent = rel_normalize_text(m.group(1))
        if not parent or rel_is_similar_name(parent, company_name):
            continue
        score = rel_source_score(url, homepage_url=homepage_url)
        return {"name": parent, "note": f"本文から抽出: {url}", "url": url, "source_type": "page_content", "confidence": rel_confidence_from_score(score, strong_signal=True)}
    if strong_signal:
        names = rel_extract_company_names(text)
        candidates = [n for n in names if not rel_is_similar_name(n, company_name)]
        candidates.sort(key=lambda n: ("HD" not in n and "ホールディングス" not in n, len(n)))
        if candidates:
            parent = rel_normalize_text(candidates[0])
            return {"name": parent, "note": f"本文の強い資本シグナルから候補抽出: {url}", "url": url, "source_type": "page_content_fallback", "confidence": rel_confidence_from_score(rel_source_score(url, homepage_url=homepage_url), strong_signal=True)}
    return None


def rel_search_relationship_sources(company_name: str, homepage_url: str, cfg: dict, log_func=None) -> dict:
    queries = [
        f"{company_name} 子会社化", f"{company_name} 株式取得", f"{company_name} 資本提携", f"{company_name} M&A",
        f"{company_name} 親会社", f"{company_name} グループ", f"{company_name} 傘下", f"{company_name} グループの一員",
        f"{company_name} site:nikkei.com", f"{company_name} site:co.jp",
    ]
    if homepage_url:
        netloc = urllib.parse.urlparse(homepage_url).netloc
        if netloc:
            queries.append(f"{company_name} site:{netloc}")
    news_items, parents_confirmed, parents_inferred, checked_urls = [], [], [], []
    for query in queries:
        results = rel_search_web(query, cfg, log_func=log_func, max_results=8)
        for item in results:
            title, url = item.get("title", ""), normalize_duckduckgo_url(item.get("url", ""))
            snippet = item.get("snippet", "")
            if not url or "houjin.info/detail" in url:
                continue
            news_items.append({"title": title, "url": url, "source": rel_domain_label(url), "snippet": snippet, "note": "検索結果"})
            combined = f"{title} {snippet}"
            if rel_contains_strong_signal(combined):
                candidate = rel_extract_parent_candidate_from_title(company_name, title)
                if candidate and not rel_is_similar_name(candidate, company_name):
                    parents_inferred.append({"name": candidate, "note": f"検索結果タイトルから推定: {title}", "url": url, "source_type": "search_title", "confidence": rel_confidence_from_score(rel_source_score(url, homepage_url=homepage_url), strong_signal=True)})
        for item in results[:REL_MAX_RESULT_URLS_TO_OPEN]:
            url = normalize_duckduckgo_url(item.get("url", ""))
            if not url or url in checked_urls:
                continue
            checked_urls.append(url)
            text = fetch_url_text(url, max_chars=REL_MAX_PAGE_TEXT_LEN, log_func=log_func)
            if not text:
                continue
            parent = rel_extract_parent_from_page_content(company_name, url, text, homepage_url=homepage_url)
            if parent:
                parents_confirmed.append(parent)
    return {
        "news_items": rel_unique_by_url(news_items),
        "parents_confirmed": rel_unique_by_name(parents_confirmed),
        "parents_inferred": rel_unique_by_name(parents_inferred),
        "checked_urls": checked_urls,
    }


def rel_search_edinet_links(keyword: str, cfg: dict, log_func=None) -> list[dict]:
    query = f"site:disclosure.edinet-fsa.go.jp {keyword} 有価証券報告書"
    results = rel_search_web(query, cfg, log_func=log_func, max_results=8)
    edinet_links = []
    for item in results:
        title = item.get("title", "")
        url = normalize_duckduckgo_url(item.get("url", ""))
        if "edinet" in url.lower() and ("有価証券報告書" in title or "edinet" in url.lower()):
            edinet_links.append({"title": title or "EDINETリンク", "url": url})
    return rel_unique_by_url(edinet_links)


def rel_scrape_edinet(company_name: str, parent_candidates: list[dict], cfg: dict, log_func=None) -> list[dict]:
    links = rel_search_edinet_links(company_name, cfg, log_func=log_func)
    for parent in parent_candidates[:2]:
        name = parent.get("name", "")
        if name:
            links.extend(rel_search_edinet_links(name, cfg, log_func=log_func))
    return rel_unique_by_url(links)


def rel_call_openai_inference(company_name: str, official_site_info: dict, confirmed_parents: list[dict], inferred_parents: list[dict], news_info: list[dict], edinet_info: list[dict], api_key: str, model_name: str = "gpt-4.1-mini") -> dict:
    if not api_key:
        return {"parent_candidate": "推論不可", "reason": "OpenAI API Key が設定されていません。", "confidence": "低"}
    try:
        system_prompt = "あなたは日本企業の資本関係を整理するアナリストです。確認済み事実を優先し、推定と事実を混同しないでください。"
        user_prompt = f"""以下の企業情報から、親会社の可能性が高い企業を1社だけ推論してください。\n\n【企業名】\n{company_name}\n\n【公式サイト情報】\n{json.dumps(official_site_info, ensure_ascii=False, indent=2)}\n\n【確認済み親会社候補】\n{json.dumps(confirmed_parents, ensure_ascii=False, indent=2)}\n\n【推定親会社候補】\n{json.dumps(inferred_parents, ensure_ascii=False, indent=2)}\n\n【検索結果】\n{json.dumps(news_info[:20], ensure_ascii=False, indent=2)}\n\n【EDINET情報】\n{json.dumps(edinet_info, ensure_ascii=False, indent=2)}\n\n【出力ルール】\nJSON形式のみで返してください。\nキーは以下の3つだけにしてください。\n- parent_candidate\n- reason\n- confidence\n\nconfidence は「高」「中」「低」のいずれか。確認済み情報が十分ならそれを優先し、不足時のみ推定してください。"""
        raw = openai_call_text(api_key, system_prompt, user_prompt, model_name=model_name)
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if not m:
            return {"parent_candidate": "推論不可", "reason": f"OpenAI応答をJSONとして解釈できませんでした: {raw[:300]}", "confidence": "低"}
        data = json.loads(m.group(0))
        return {"parent_candidate": str(data.get("parent_candidate", "推論不可")), "reason": str(data.get("reason", "")), "confidence": str(data.get("confidence", "低"))}
    except Exception as e:
        return {"parent_candidate": "推論不可", "reason": f"OpenAI推論エラー: {e}", "confidence": "低"}


def investigate_company_relationship_data(company_name: str, company_url: str, cfg: dict, log_func=None) -> dict:
    execution_notes = []
    resolved_url = (company_url or "").strip()
    if not resolved_url:
        resolved_url = rel_guess_company_url(company_name, cfg, log_func=log_func)
        execution_notes.append(f"URL自動推定: {resolved_url if resolved_url else '失敗'}")
    else:
        execution_notes.append(f"URL指定あり: {resolved_url}")
    official_info = {"subsidiaries": [], "affiliates": [], "official_pages_checked": []}
    if resolved_url:
        official_info = rel_scrape_official_site(resolved_url, log_func=log_func)
        execution_notes.append(f"公式サイト確認ページ数: {len(official_info.get('official_pages_checked', []))}")
    else:
        execution_notes.append("公式サイト解析は未実施")
    search_info = rel_search_relationship_sources(company_name, resolved_url, cfg, log_func=log_func)
    confirmed_parents = search_info.get("parents_confirmed", [])
    inferred_parents = search_info.get("parents_inferred", [])
    news_items = search_info.get("news_items", [])
    checked_urls = search_info.get("checked_urls", [])
    execution_notes.append(f"検索結果URL確認数: {len(checked_urls)}")
    execution_notes.append(f"確認済み親会社候補数: {len(confirmed_parents)}")
    execution_notes.append(f"推定親会社候補数: {len(inferred_parents)}")
    edinet_links = rel_scrape_edinet(company_name, confirmed_parents or inferred_parents, cfg, log_func=log_func)
    execution_notes.append(f"EDINETリンク数: {len(edinet_links)}")
    openai_inference = rel_call_openai_inference(
        company_name=company_name,
        official_site_info=official_info,
        confirmed_parents=confirmed_parents,
        inferred_parents=inferred_parents,
        news_info=news_items,
        edinet_info=edinet_links,
        api_key=cfg.get("openai_api_key", ""),
        model_name=cfg.get("openai_model", "gpt-4.1-mini"),
    )
    return {
        "company_name": company_name,
        "company_url": resolved_url,
        "confirmed_parents": confirmed_parents,
        "inferred_parents": inferred_parents,
        "subsidiaries": official_info.get("subsidiaries", []),
        "affiliates": official_info.get("affiliates", []),
        "official_pages_checked": official_info.get("official_pages_checked", []),
        "edinet_links": edinet_links,
        "news_items": news_items,
        "checked_urls": checked_urls,
        "openai_inference": openai_inference,
        "execution_notes": execution_notes,
    }


def rel_render_relation_item(item: dict) -> list[str]:
    lines = [f"- {item.get('name', '')}"]
    if item.get("confidence"):
        lines.append(f"  - 確度: {item.get('confidence', '')}")
    if item.get("note"):
        lines.append(f"  - 根拠: {item.get('note', '')}")
    if item.get("url"):
        lines.append(f"  - URL: {item.get('url', '')}")
    if item.get("source_type"):
        lines.append(f"  - 種別: {item.get('source_type', '')}")
    return lines


def build_relationship_markdown_section(result: dict) -> str:
    company_name = result.get("company_name", "")
    company_url = result.get("company_url", "")
    confirmed_parents = result.get("confirmed_parents", [])
    inferred_parents = result.get("inferred_parents", [])
    subsidiaries = result.get("subsidiaries", [])
    affiliates = result.get("affiliates", [])
    edinet_links = result.get("edinet_links", [])
    news_items = result.get("news_items", [])
    checked_urls = result.get("checked_urls", [])
    ai = result.get("openai_inference", {})
    execution_notes = result.get("execution_notes", [])
    lines = []
    lines.append("\n\n---\n")
    lines.append("# 企業資本関係調査レポート")
    lines.append("")
    lines.append("## 調査対象企業情報")
    lines.append(f"- 企業名: {company_name}")
    lines.append(f"- 企業URL: {company_url if company_url else '自動推定不可'}")
    lines.append("")
    lines.append("## 親会社候補（確認済み事実）")
    if confirmed_parents:
        for item in confirmed_parents:
            lines.extend(rel_render_relation_item(item))
    else:
        lines.append("- なし")
    lines.append("")
    lines.append("## 親会社候補（推定情報）")
    if inferred_parents:
        for item in inferred_parents:
            lines.extend(rel_render_relation_item(item))
    else:
        lines.append("- なし")
    lines.append("")
    lines.append("## 子会社一覧（公式サイト確認結果）")
    if subsidiaries:
        for item in subsidiaries:
            lines.extend(rel_render_relation_item(item))
    else:
        lines.append("- なし")
    lines.append("")
    lines.append("## 関係会社一覧（公式サイト確認結果）")
    if affiliates:
        for item in affiliates:
            lines.extend(rel_render_relation_item(item))
    else:
        lines.append("- なし")
    lines.append("")
    lines.append("## EDINET 有価証券報告書リンク")
    if edinet_links:
        for link in edinet_links:
            lines.append(f"- [{link.get('title', 'EDINETリンク')}]({link.get('url', '')})")
    else:
        lines.append("- なし")
    lines.append("")
    lines.append("## OpenAIによる最終補完推論")
    lines.append(f"- 親会社候補: {ai.get('parent_candidate', '推論不可')}")
    lines.append(f"- 根拠: {ai.get('reason', '')}")
    lines.append(f"- 確度: {ai.get('confidence', '低')}")
    lines.append("")
    lines.append("## 検索・確認したURL")
    if checked_urls:
        for url in checked_urls[:30]:
            lines.append(f"- {url}")
    else:
        lines.append("- なし")
    lines.append("")
    lines.append("## 参考検索結果")
    if news_items:
        for item in news_items[:20]:
            lines.append(f"- [{item.get('title', '')}]({item.get('url', '')})")
            lines.append(f"  - ソース: {item.get('source', '')}")
            if item.get("snippet"):
                lines.append(f"  - 要約: {item.get('snippet', '')}")
    else:
        lines.append("- なし")
    lines.append("")
    lines.append("## 実行メモ")
    for note in execution_notes:
        lines.append(f"- {note}")
    lines.append("")
    lines.append("## 調査フレームワーク")
    lines.append("- 公式サイトのグループ情報を確認")
    lines.append("- ニュース・企業DB・公式IRから資本関係情報を確認")
    lines.append("- EDINETで有価証券報告書リンクを確認")
    lines.append("- OpenAI推論により不足情報を補完")
    lines.append("- レポートでは確認済み事実と推定情報を分離表示")
    lines.append("")
    return "\n".join(lines)

def generate_report(company_name, homepage_url, cfg, gmark_excel_path, workplace_excel_path, nta_csv_path, log_func):
    plan, sources = collect_sources(company_name, homepage_url, cfg["openai_api_key"], cfg, log_func=log_func, model_name=cfg.get("openai_model", "gpt-4.1-mini"))
    if not sources:
        raise RuntimeError("有効なソースが取得できませんでした。企業名やURLを確認してください。")
    report_md = build_markdown_report(company_name, homepage_url, plan, sources, cfg["openai_api_key"], log_func=log_func, model_name=cfg.get("openai_model", "gpt-4.1-mini"))
    identity_result = analyze_company_identity(company_name, homepage_url, sources, cfg["openai_api_key"], log_func=log_func, model_name=cfg.get("openai_model", "gpt-4.1-mini"))
    nta_raw_result = search_corporate_number_from_nta(identity_result.get("official_company_name", company_name), identity_result.get("head_office_address", ""), log_func=log_func, csv_path=nta_csv_path)
    corporate_number_result = choose_best_corporate_number_candidate(identity_result.get("official_company_name", company_name), identity_result.get("head_office_address", ""), nta_raw_result)
    permit_sources = search_general_freight_permit_by_corporate_number(identity_result.get("official_company_name", company_name), (corporate_number_result.get("best") or {}).get("corporate_number", ""), cfg, log_func=log_func)
    permit_result = evaluate_general_freight_permit(identity_result.get("official_company_name", company_name), permit_sources, log_func=log_func)
    gmark_result = search_company_in_excel(gmark_excel_path, company_name, log_func=log_func)
    workplace_result = search_company_in_excel(workplace_excel_path, company_name, log_func=log_func)
#    appended_section = build_additional_registry_section(company_name, identity_result, corporate_number_result, permit_result, gmark_result, workplace_result)
    appended_section = ""
    relationship_result = investigate_company_relationship_data(company_name, homepage_url, cfg, log_func=log_func)
    relationship_section = build_relationship_markdown_section(relationship_result)
    return report_md + appended_section + relationship_section


def main():
    st.set_page_config(page_title="企業情報調査ツール Trial版", layout="wide",initial_sidebar_state="collapsed")
    cfg = load_config()

    secret_openai_api_key = get_streamlit_secret("OPENAI_API_KEY", "")
    secret_serpapi_key = get_streamlit_secret("SERPAPI_KEY", "")

    st.title("企業情報調査ツール(Trial版)")
    st.caption("設定を変更する場合には、サイドバーを開いて設定してください")
#    st.success("画面初期化OK")

    if "logs" not in st.session_state:
        st.session_state.logs = []
    if "report_md" not in st.session_state:
        st.session_state.report_md = ""
    if "report_docx_bytes" not in st.session_state:
        st.session_state.report_docx_bytes = b""

    def log_func(msg):
        ts = datetime.now().strftime("%H:%M:%S")
        st.session_state.logs.append(f"[{ts}] {msg}")

    try:
        with st.sidebar:
            st.subheader("API設定")
            st.text_input(
                "OpenAI API Key",
                value="Secretsから読込済み" if secret_openai_api_key else "Secrets未設定",
                type="password",
                disabled=True,
            )
            openai_model = st.text_input("OpenAI モデル名", value=cfg.get("openai_model", "gpt-4.1-mini"))
            st.text_input(
                "SerpAPI Key",
                value="Secretsから読込済み" if secret_serpapi_key else "Secrets未設定",
                type="password",
                disabled=True,
            )
            use_serpapi = st.checkbox("SerpAPI を使用", value=cfg.get("use_serpapi", True))
            email_domain = st.text_input("メールドメイン（任意）", value=cfg.get("email_domain", ""))

            if st.button("設定を保存"):
                cfg.update({
                    "openai_model": openai_model,
                    "use_serpapi": bool(use_serpapi and secret_serpapi_key),
                    "email_domain": email_domain,
                })
                cfg.pop("openai_api_key", None)
                cfg.pop("serpapi_key", None)
                save_config(cfg)
                st.success("設定を保存しました。")

            st.markdown("### 補助ファイル")
            gmark_excel_path_text = st.text_input("Gマーク認証ExcelのWindowsパス", value=cfg.get("gmark_excel_path", ""), help=r"例: C:\Data\gmark.xlsx")
            gmark_excel_upload = st.file_uploader("Gマーク認証Excelをアップロード", type=["xlsx", "xlsm", "xltx", "xltm"], key="gmark")
            workplace_excel_path_text = st.text_input("働きやすい職場認証ExcelのWindowsパス", value=cfg.get("workplace_excel_path", ""), help=r"例: C:\Data\workplace.xlsx")
            workplace_excel_upload = st.file_uploader("働きやすい職場認証Excelをアップロード", type=["xlsx", "xlsm", "xltx", "xltm"], key="workplace")
            nta_csv_path_text = st.text_input("法人番号CSV（Shift-JIS）のWindowsパス", value=cfg.get("nta_csv_path", ""), help=r"例: C:\Data\houjin.csv")
            nta_csv_upload = st.file_uploader("法人番号CSVをアップロード", type=["csv"], key="nta")
            output_docx_path = st.text_input("出力Word保存先（Windowsローカル保存用、任意）", value=cfg.get("last_output_docx_path", ""), help=r"例: C:\Reports\company_report.docx")


        with st.form("main_form"):
            col1, col2 = st.columns(2)
            with col1:
                company_name = st.text_input("企業名", value=cfg.get("last_company_name", ""))
            with col2:
                homepage_url = st.text_input("ホームページURL（必須）", value=cfg.get("last_homepage_url", ""))

            submitted = st.form_submit_button("調査開始", type="primary")

        if submitted:
            st.session_state.logs = []
            st.session_state.report_md = ""
            st.session_state.report_docx_bytes = b""

            errors = []
            if not company_name.strip():
                errors.append("企業名を入力してください。")
            if not homepage_url.strip():
                errors.append("ホームページURLを入力してください。")
            if not secret_openai_api_key.strip():
                errors.append("OpenAI API Key を設定してください。")

            if errors:
                for err in errors:
                    st.error(err)
            else:
                cfg.update({
                    "openai_api_key": secret_openai_api_key.strip(),
                    "openai_model": openai_model.strip() or "gpt-4.1-mini",
                    "serpapi_key": secret_serpapi_key.strip(),
                    "use_serpapi": bool(use_serpapi and secret_serpapi_key.strip()),
                    "email_domain": email_domain.strip(),
                    "last_company_name": company_name.strip(),
                    "last_homepage_url": homepage_url.strip(),
                    "gmark_excel_path": gmark_excel_path_text.strip(),
                    "workplace_excel_path": workplace_excel_path_text.strip(),
                    "nta_csv_path": nta_csv_path_text.strip(),
                    "last_output_docx_path": output_docx_path.strip(),
                })
                save_cfg = dict(cfg)
                save_cfg.pop("openai_api_key", None)
                save_cfg.pop("serpapi_key", None)
                save_config(save_cfg)

                gmark_excel_path = resolve_input_path("Gマーク", gmark_excel_path_text.strip(), gmark_excel_upload)
                workplace_excel_path = resolve_input_path("働きやすい職場", workplace_excel_path_text.strip(), workplace_excel_upload)
                nta_csv_path = resolve_input_path("法人番号CSV", nta_csv_path_text.strip(), nta_csv_upload)

                try:
                    with st.spinner("調査中です..."):
                        report_md = generate_report(company_name.strip(), homepage_url.strip(), cfg, gmark_excel_path, workplace_excel_path, nta_csv_path, log_func)
                        st.session_state.report_md = report_md
                        report_docx_bytes = markdown_to_docx_bytes(report_md)
                        st.session_state.report_docx_bytes = report_docx_bytes

                        saved_path = ""
                        if output_docx_path.strip():
                            try:
                                out = Path(output_docx_path.strip())
                                out.parent.mkdir(parents=True, exist_ok=True)
                                out.write_bytes(report_docx_bytes)
                                saved_path = str(out)
                            except Exception as e:
                                log_func(f"[WARN] 指定Wordパスへの保存に失敗: {e}")

                        if saved_path:
                            st.success(f"調査完了。Wordファイルをローカル保存しました: {saved_path}")
                        else:
                            st.success("調査完了。下のダウンロードボタンからWordファイルを保存できます。")
                except Exception as e:
                    import traceback
                    st.error(f"調査中にエラーが発生しました: {e}")
                    st.code(traceback.format_exc(), language="python")

        st.markdown("### レポート")
        if st.session_state.report_md:
            cdl1, cdl2 = st.columns(2)
            with cdl1:
                st.download_button(
                    "Wordをダウンロード",
                    data=st.session_state.report_docx_bytes,
                    file_name=f"company_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            with cdl2:
                st.download_button(
                    "Markdownをダウンロード",
                    data=st.session_state.report_md.encode("utf-8"),
                    file_name=f"company_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown",
                )
            st.markdown("#### ブラウザ表示（Markdown整形）")
            st.markdown(st.session_state.report_md)
            with st.expander("Markdown原文を表示"):
                st.text_area("生成レポート原文", value=st.session_state.report_md, height=520)
        else:
            st.info("まだレポートは生成されていません。")

    except Exception as e:
        import traceback
        st.error(f"初期描画でエラーが発生しました: {e}")
        st.code(traceback.format_exc(), language="python")


if __name__ == "__main__":
    main()
